#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
@author: 'Guohan'
@file: getEnglish_txt_from_docx.py
@time: 2/6/2025 下午 11:32
@functions：从一个Word文件中，提取仅英文的内容，追加到box.docx文件里，注意如果一段中文文字一有一些小于10个英文字母，那么也可以忽略

"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


# 正则匹配纯英文和常见标点符号
def extract_english_only(text):
    return ''.join(re.findall(r'[a-zA-Z0-9\s,.!?;:\(\)\'"]+', text))


# 复制运行格式
def copy_run_format(dest_run, src_run):
    dest_run.bold = src_run.bold
    dest_run.italic = src_run.italic
    dest_run.underline = src_run.underline
    dest_run.font.name = src_run.font.name
    if src_run.font.size:
        dest_run.font.size = src_run.font.size
    if src_run.font.color and src_run.font.color.rgb:
        dest_run.font.color.rgb = src_run.font.color.rgb
    else:
        dest_run.font.color.rgb = RGBColor(0, 0, 0)


# 插入分隔线
def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('-' * 50)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)


# 判断是否是句尾（以句号、问号、感叹号结尾）
def is_sentence_end(text):
    return re.search(r'[.!?]\s*$', text.strip())


# 判断是否是小写字母开头
def starts_with_lowercase(text):
    stripped = text.lstrip()
    return stripped and stripped[0].islower()


# 主函数（优化了段落合并逻辑）
def process_and_append(source_path, target_path):
    source_doc = Document(source_path)
    target_doc = Document(target_path) if os.path.exists(target_path) else Document()

    # 添加分隔线
    if len(target_doc.paragraphs) > 0:
        add_separator(target_doc)

    # 存储所有提取出的段落内容（run 级别合并后的文本）
    extracted_paragraphs = []

    for para in source_doc.paragraphs:
        new_text = ''
        for run in para.runs:
            clean_text = extract_english_only(run.text)
            new_text += clean_text
        if new_text.strip():
            extracted_paragraphs.append(new_text.strip())

    # 合并段落逻辑
    merged_paragraphs = []
    i = 0
    while i < len(extracted_paragraphs):
        current = extracted_paragraphs[i]
        # 如果不是以句号结尾，并且下一段以小写字母开头
        while i + 1 < len(extracted_paragraphs) and not is_sentence_end(current) and \
                starts_with_lowercase(extracted_paragraphs[i + 1]):
            next_para = extracted_paragraphs[i + 1]
            current += ' ' + next_para.strip()
            i += 1
        merged_paragraphs.append(current)
        i += 1

    # 将合并后的段落写入文档
    for para_text in merged_paragraphs:
        new_para = target_doc.add_paragraph()
        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_run = new_para.add_run(para_text)
        # 如需保留每个 run 的详细格式，请参考原始完整版代码

    # 保存目标文档
    target_doc.save(target_path)
    print(f"处理完成，内容已追加至 {target_path}")


# ========== 用户交互入口 ==========
if __name__ == '__main__':
    print("📘 英文内容提取器 -开始")
    source_file = "Purpose-Profit.docx"
    target_file = "box.docx"

    print("\n🔄 开始处理，请稍等...\n")
    process_and_append(source_file, target_file)
