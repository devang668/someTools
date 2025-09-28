#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
@author: 'Guohan'
@file: getEnglish_txt_from_docx_pro.py
@time: 3/6/2025 上午 12:44
@functions：please note ..
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


# 正则匹配纯英文和常见标点符号
def extract_english_only(text):
    return ''.join(re.findall(r'[a-zA-Z0-9\s,.!?;:\(\)\'"]+', text))


# 复制运行格式
def copy_run_format(dest_run, src_run):
    dest_run.bold = src_run.bold
    dest_run.italic = src_run.italic
    dest_run.underline = src_run.underline
    dest_run.font.name = src_run.font.name
    if src_run.font.size:
        dest_run.font.size = src_run.font.size
    if src_run.font.color and src_run.font.color.rgb:
        dest_run.font.color.rgb = src_run.font.color.rgb
    else:
        dest_run.font.color.rgb = RGBColor(0, 0, 0)


# 插入分隔线
def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('-' * 50)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)


# 判断是否是句尾（以句号、问号、感叹号结尾）
def is_sentence_end(text):
    return re.search(r'[.!?]\s*$', text.strip())


# 判断是否是小写字母开头
def starts_with_lowercase(text):
    stripped = text.lstrip()
    return stripped and stripped[0].islower()


# 主函数（优化了段落合并逻辑）

def process_and_append(source_path, target_path):
    source_doc = Document(source_path)
    target_doc = Document(target_path) if os.path.exists(target_path) else Document()

    # 添加分隔线
    if len(target_doc.paragraphs) > 0:
        add_separator(target_doc)

    for para in source_doc.paragraphs:
        new_para = target_doc.add_paragraph()

        # 复制段落样式（如标题级别）
        if para.style and para.style.name:
            try:
                new_para.style = target_doc.styles[para.style.name]
            except KeyError:
                # 如果目标文档没有这个样式，则使用默认段落样式
                new_para.style = target_doc.styles['Normal']

        # 处理每个 run
        for run in para.runs:
            clean_text = extract_english_only(run.text)
            if not clean_text:
                continue
            new_run = new_para.add_run(clean_text)
            copy_run_format(new_run, run)

        # 如果整个段落为空（全是中文），则删除这个空段落
        if not new_para.text.strip():
            p = new_para._element
            p.getparent().remove(p)
            continue

    target_doc.save(target_path)
    print(f"处理完成，内容已追加至 {target_path}")

# ========== 用户交互入口 ==========
if __name__ == '__main__':
    print("📘 英文内容提取器 -开始")
    source_file = "Purpose-Profit.docx"
    target_file = "box.docx"

    print("\n🔄 开始处理，请稍等...\n")
    process_and_append(source_file, target_file)